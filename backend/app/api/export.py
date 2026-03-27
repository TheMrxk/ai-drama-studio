"""
Export API Routes
"""
from flask import Blueprint, request, jsonify, send_file
from flask_jwt_extended import jwt_required, get_jwt_identity
from app.models import db, Project, Version
import io
import os

bp = Blueprint('export', __name__, url_prefix='/export')


@bp.route('/<project_id>/txt', methods=['GET'])
@jwt_required()
def export_txt(project_id):
    """Export script as plain text"""
    current_user_id = get_jwt_identity()

    # Get project
    project = Project.query.filter_by(id=project_id, user_id=current_user_id).first()
    if not project:
        return jsonify({'error': 'Project not found'}), 404

    # Get latest version
    latest_version = Version.query.filter_by(project_id=project_id).order_by(Version.version.desc()).first()
    content = latest_version.content if latest_version else generate_default_content(project)

    # Create text file
    buffer = io.BytesIO()
    buffer.write(content.encode('utf-8'))
    buffer.seek(0)

    return send_file(
        buffer,
        mimetype='text/plain',
        as_attachment=True,
        download_name=f'{project.name}.txt'
    )


@bp.route('/<project_id>/pdf', methods=['GET'])
@jwt_required()
def export_pdf(project_id):
    """Export script as PDF"""
    current_user_id = get_jwt_identity()

    # Get project
    project = Project.query.filter_by(id=project_id, user_id=current_user_id).first()
    if not project:
        return jsonify({'error': 'Project not found'}), 404

    # Get latest version
    latest_version = Version.query.filter_by(project_id=project_id).order_by(Version.version.desc()).first()
    content = latest_version.content if latest_version else generate_default_content(project)

    # Create PDF using reportlab
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.enums import TA_LEFT
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.lib import colors

    # Try to register Chinese font
    font_paths = [
        '/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc',
        '/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc',
        '/System/Library/Fonts/PingFang.ttc',
        'C:\\Windows\\Fonts\\simhei.ttf',
    ]

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, leftMargin=2*cm, rightMargin=2*cm, topMargin=2*cm, bottomMargin=2*cm)
    styles = getSampleStyleSheet()

    # Custom style for Chinese
    try:
        for font_path in font_paths:
            if os.path.exists(font_path):
                pdfmetrics.registerFont(TTFont('Chinese', font_path))
                break
        title_style = ParagraphStyle(name='CustomTitle', parent=styles['Heading1'], fontSize=18, leading=24, alignment=TA_LEFT)
        normal_style = ParagraphStyle(name='CustomNormal', parent=styles['Normal'], fontSize=11, leading=16, fontName='Chinese')
    except:
        title_style = ParagraphStyle(name='CustomTitle', parent=styles['Heading1'], fontSize=18, leading=24)
        normal_style = ParagraphStyle(name='CustomNormal', parent=styles['Normal'], fontSize=11, leading=16)

    # Build PDF content
    story = []

    # Title
    story.append(Paragraph(f"《{project.name}》", title_style))
    story.append(Spacer(1, 0.5*cm))

    # Meta info
    meta_info = f"风格：{project.style} | 集数：{project.episodes} 集"
    story.append(Paragraph(meta_info, normal_style))
    story.append(Spacer(1, 0.3*cm))

    # Setting
    story.append(Paragraph(f"故事设定：{project.setting}", normal_style))
    story.append(Spacer(1, 0.3*cm))

    # Characters
    story.append(Paragraph(f"主角特征：{project.characters}", normal_style))
    story.append(Spacer(1, 0.3*cm))

    # Plot
    story.append(Paragraph(f"剧情脉络：{project.plot}", normal_style))
    story.append(Spacer(1, 0.5*cm))

    # Script content
    story.append(Paragraph("=" * 50, normal_style))
    story.append(Spacer(1, 0.3*cm))

    # Parse and add script content line by line
    for line in content.split('\n'):
        if line.strip():
            # Escape special characters for reportlab
            escaped_line = line.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            story.append(Paragraph(escaped_line, normal_style))

    story.append(Spacer(1, 0.3*cm))

    # Build PDF
    doc.build(story)
    buffer.seek(0)

    return send_file(
        buffer,
        mimetype='application/pdf',
        as_attachment=True,
        download_name=f'{project.name}.pdf'
    )


@bp.route('/<project_id>/docx', methods=['GET'])
@jwt_required()
def export_docx(project_id):
    """Export script as Word document"""
    current_user_id = get_jwt_identity()

    # Get project
    project = Project.query.filter_by(id=project_id, user_id=current_user_id).first()
    if not project:
        return jsonify({'error': 'Project not found'}), 404

    # Get latest version
    latest_version = Version.query.filter_by(project_id=project_id).order_by(Version.version.desc()).first()
    content = latest_version.content if latest_version else generate_default_content(project)

    # Create Word document using python-docx
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Title
    title = doc.add_heading(f"《{project.name}》", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Meta info
    doc.add_paragraph(f"风格：{project.style} | 集数：{project.episodes} 集")

    # Project details
    doc.add_heading('故事设定', level=2)
    doc.add_paragraph(project.setting)

    doc.add_heading('主角特征', level=2)
    doc.add_paragraph(project.characters)

    doc.add_heading('剧情脉络', level=2)
    doc.add_paragraph(project.plot)

    doc.add_heading('最终结局', level=2)
    doc.add_paragraph(project.ending)

    doc.add_heading('剧本内容', level=2)

    # Add script content
    for line in content.split('\n'):
        if line.strip():
            doc.add_paragraph(line)

    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return send_file(
        buffer,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name=f'{project.name}.docx'
    )


def generate_default_content(project):
    """Generate default content when no version exists"""
    return f"""# 《{project.name}》

## 基本信息
- 风格：{project.style}
- 集数：{project.episodes} 集

## 故事设定
{project.setting}

## 主角特征
{project.characters}

## 剧情脉络
{project.plot}

## 最终结局
{project.ending}

---

**剧本内容待生成**
"""
